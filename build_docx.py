#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Build a printable DOCX from merged Markdown.

Usage:
  python build_docx.py <md_file> <docx_file> [volume_title]
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


BLUE_DARK = RGBColor(0x1F, 0x4E, 0x79)
BLUE_MID = RGBColor(0x2E, 0x75, 0xB6)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY_BG = "F2F2F2"

FONT_BODY = "Microsoft YaHei"
FONT_CODE = "Microsoft YaHei"
FONT_SIZE_TITLE = Pt(22)
FONT_SIZE_SUBTITLE = Pt(14)
FONT_SIZE_H1 = Pt(14)
FONT_SIZE_H2 = Pt(11.5)
FONT_SIZE_H3 = Pt(10)
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_CODE = Pt(9)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_run_font(run, name, size, color=BLACK, bold=False):
    run.bold = bold
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_field_run(paragraph, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field_run(paragraph, "PAGE")


def add_toc(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field_run(paragraph, r'TOC \o "1-2" \h \z \u')


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = FONT_SIZE_BODY
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = FONT_BODY
    heading1.font.size = FONT_SIZE_H1
    heading1.font.bold = True
    heading1.font.color.rgb = BLUE_DARK
    heading1.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = FONT_BODY
    heading2.font.size = FONT_SIZE_H2
    heading2.font.bold = True
    heading2.font.color.rgb = BLUE_MID
    heading2.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = FONT_BODY
    heading3.font.size = FONT_SIZE_H3
    heading3.font.bold = True
    heading3.font.color.rgb = BLACK
    heading3.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def add_cover_page(doc, volume_title):
    section = doc.sections[0]
    configure_section(section)
    section.different_first_page_header_footer = True

    # Split title by ：or —— into main title + subtitle
    delimiter = None
    for d in ("：", "——"):
        if d in volume_title:
            delimiter = d
            break
    if delimiter:
        parts = volume_title.split(delimiter, 1)
        main_title = parts[0].strip()
        subtitle = parts[1].strip() if len(parts) > 1 else ""
    else:
        main_title = volume_title
        subtitle = ""

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(main_title)
    set_run_font(run, FONT_BODY, FONT_SIZE_TITLE, bold=True)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(subtitle)
        set_run_font(run, FONT_BODY, FONT_SIZE_SUBTITLE, color=BLUE_DARK, bold=True)


def add_toc_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    section.different_first_page_header_footer = False
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_number(footer_para)

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("目录")
    set_run_font(run, FONT_BODY, FONT_SIZE_H1, color=BLUE_DARK, bold=True)

    toc_paragraph = doc.add_paragraph()
    toc_paragraph.paragraph_format.line_spacing = 1.0
    add_toc(toc_paragraph)

    doc.add_page_break()


def add_heading(doc, text, level):
    style_name = f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing = 1.0
    if level == 1:
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, FONT_BODY, FONT_SIZE_H1, color=BLUE_DARK, bold=True)
    elif level == 2:
        set_run_font(run, FONT_BODY, FONT_SIZE_H2, color=BLUE_MID, bold=True)
    else:
        set_run_font(run, FONT_BODY, FONT_SIZE_H3, bold=True)


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, FONT_BODY, FONT_SIZE_BODY, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run, FONT_BODY, FONT_SIZE_BODY)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{GRAY_BG}"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    set_run_font(run, FONT_CODE, FONT_SIZE_CODE)


def add_translation(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    label = p.add_run("中文译文：")
    set_run_font(label, FONT_BODY, FONT_SIZE_BODY, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, FONT_BODY, FONT_SIZE_BODY)


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    border = parse_xml(
        '<w:pBdr {}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>'.format(
            nsdecls("w")
        )
    )
    p._element.get_or_add_pPr().append(border)


def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• ")
    set_run_font(run, FONT_BODY, FONT_SIZE_BODY, bold=True)
    add_body_fragment(p, text)


def add_body_fragment(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, FONT_BODY, FONT_SIZE_BODY, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, FONT_BODY, FONT_SIZE_BODY)


def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    in_code_block = False
    code_buffer = []

    separator_pattern = re.compile(r"^---\s*$")
    list_pattern = re.compile(r"^-\s+(.+)")
    numbered_pattern = re.compile(r"^\d+\.\s+(.+)")
    first_h1_skipped = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if separator_pattern.match(stripped):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            add_separator(doc)
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            if first_h1_skipped:
                add_heading(doc, line[2:].strip(), 1)
            else:
                first_h1_skipped = True
            i += 1
            continue

        if line.startswith("## "):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue

        if line.startswith("### "):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue

        if line.startswith("> ") or line.startswith("    "):
            code_text = line[2:] if line.startswith("> ") else line[4:]
            if code_text.startswith("**中文译文：**"):
                if in_code_block:
                    add_code_block(doc, "\n".join(code_buffer))
                    code_buffer = []
                    in_code_block = False
                add_translation(doc, code_text[len("**中文译文：**") :].strip())
            else:
                in_code_block = True
                code_buffer.append(code_text)
            i += 1
            continue

        if in_code_block and not stripped:
            add_code_block(doc, "\n".join(code_buffer))
            code_buffer = []
            in_code_block = False
            i += 1
            continue

        match = list_pattern.match(line)
        if match:
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            add_list_item(doc, match.group(1))
            i += 1
            continue

        match = numbered_pattern.match(stripped)
        if match:
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            add_list_item(doc, match.group(1))
            i += 1
            continue

        if not stripped:
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            add_code_block(doc, "\n".join(code_buffer))
            code_buffer = []
            in_code_block = False

        add_body(doc, line)
        i += 1

    if in_code_block and code_buffer:
        add_code_block(doc, "\n".join(code_buffer))


def main():
    if len(sys.argv) < 3:
        print("Usage: python build_docx.py <md_file> <docx_file> [series_title] [volume_title]")
        sys.exit(1)

    md_file = Path(sys.argv[1])
    docx_file = Path(sys.argv[2])
    volume_title = sys.argv[3] if len(sys.argv) >= 4 else md_file.stem

    text = md_file.read_text(encoding="utf-8")

    doc = Document()
    setup_styles(doc)
    add_cover_page(doc, volume_title)
    add_toc_section(doc)
    parse_markdown(doc, text)

    docx_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_file))
    print(f"DOCX saved to: {docx_file}")


if __name__ == "__main__":
    main()
